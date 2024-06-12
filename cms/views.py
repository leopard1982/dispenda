import os
from django.conf import settings
from django.shortcuts import render, HttpResponseRedirect, HttpResponse
import docx
from docx.shared import Pt,Cm
from cms.forms import inputGolongan, inputJabatan, inputPegawai, inputPengguna, inputSuratTugas
from cms.forms import inputMasterDasarST, inputConfig, inputPesertaTugas, inputDasarSuratTugas
from surat_tugas.models import MasterGolongan,MasterJabatan,MasterPegawai, Pengguna, Logging
from surat_tugas.models import TrxSuratTugas, ConfigDispenda, MasterDasarST, ST_DasarTugas, ST_Peserta
from django.core.paginator import Paginator
from django.contrib.auth import authenticate,login,logout
from django.db.models import Q
import uuid

def handler404(request):
	return render(request,'404.html')

def addLogging(username,grouping,message):
	loggingnya=Logging.objects.create(
		user=Pengguna.objects.get(username=username),
		transaction=message,
		grouping=grouping
	)
	return loggingnya

def getPendingSurat():
	return TrxSuratTugas.objects.all().filter(submit=False)


def dashboard(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	
	jml_pending_surattugas = TrxSuratTugas.objects.all().filter(submit=False).count()
	jml_sukses_surattugas = TrxSuratTugas.objects.all().filter(submit=True).count()
	jml_master_pegawai = MasterPegawai.objects.all()
	configdispenda = ConfigDispenda.objects.all()
	if configdispenda.count()==0:
		configdispenda=None
	else:
		configdispenda=configdispenda[0]

	context={
		
		'username': request.user.username,
		'jml_pending_surattugas': jml_pending_surattugas,
		'jml_sukses_surattugas': jml_sukses_surattugas,
		'jml_master_pegawai':jml_master_pegawai,
		'kepala':configdispenda,
		'pending_surat':getPendingSurat()
	}
	return render(request,'master/dashboard.html',context)

def addGolongan(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	if request.method=="POST":
		kode_golongan = request.POST['kode_golongan']
		keterangan = request.POST['keterangan']

		try:
			golongan = MasterGolongan.objects.create(
				kode_golongan=kode_golongan,
				keterangan=keterangan,
				updatedBy=request.user.username
			)
			golongan.save()
			addLogging(request.user.username,"master_golongan",f"berhasil - create kode: {kode_golongan} - {keterangan}" )
			request.session['status']='Penambahan Master Golongan Berhasil!'
		except:
			request.session['status']="Penambahan Master Golongan Gagal! Data Pernah Dibuat!"
			addLogging(request.user.username,"master_golongan",f"gagal[data sudah ada] - create kode: {kode_golongan} - {keterangan}" )
		return HttpResponseRedirect('/master/gol/dis/')
	
	context={
		'forms':inputGolongan,
		'menuname':'Menambah Golongan',
		'pathway':'Master Golongan - Menambah Golongan',
		'username': request.user.username,
		'status':request.session['status'],
		'pending_surat':getPendingSurat()
	}
	return render(request,'master/create_golongan.html',context)

def displayGolongan(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		page=request.GET['page']
	except:
		page=1
	if page==None:
		page=1
	listGolongan = MasterGolongan.objects.all().order_by('-updatedAt')
	p=Paginator(listGolongan,5)
	
	try:
		halaman = p.page(page)
	except:
		halaman = p.page(1)
	context={
		'lists':halaman,
		'paginator':p,
		'menuname':'Daftar Golongan',
		'pathway':'Master Golongan - Daftar Golongan',
		'username':request.user.username,
		'mywebsite': '/master/gol/dis/',
		'status':request.session['status'],
		'pending_surat':getPendingSurat()
	}
	request.session['status']=""
	return render(request,'master/display_golongan.html',context)


def addJabatan(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	if request.method=="POST":
		kode_jabatan = request.POST['kode_jabatan']
		keterangan = request.POST['keterangan']

		try:
			jabatan = MasterJabatan.objects.create(
				kode_jabatan=kode_jabatan,
				keterangan=keterangan,
				updatedBy=request.user.username
			)
			jabatan.save()
			addLogging(request.user.username,"master_jabatan",f"berhasil - create kode: {kode_jabatan} - {keterangan}" )
			request.session['status']="Penambahan Master Jabatan Berhasil!"
		except:
			request.session['status']="Penambahan Master Jabatan Gagal! Data pernah dibuat!"
			addLogging(request.user.username,"master_golongan",f"gagal[data sudah ada] - create kode: {kode_golongan} - {keterangan}" )
		return HttpResponseRedirect('/master/jab/dis/')

	context={
		'forms':inputJabatan,
		'menuname':'Menambah Jabatan',
		'pathway':'Master Jabatan - Menambah Jabatan',
		'username':request.user.username,
		'pending_surat':getPendingSurat()
	}
	return render(request,'master/create_jabatan.html',context)

def displayJabatan(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		page=request.GET['page']
	except:
		page=1
	if page==None:
		page=1

	listJabatan = MasterJabatan.objects.all().order_by('-updatedAt')
	
	p=Paginator(listJabatan,5)
	
	
	try:
		halaman = p.page(page)
	except:
		halaman = p.page(1)
	
	context={
		'lists':halaman,
		'paginator':p,
		'menuname':'Daftar Jabatan',
		'pathway':'Master Jabatan - Daftar Jabatan',
		'username':request.user.username,
		'mywebsite': '/master/jab/dis/',
		'status':request.session['status'],
		'pending_surat':getPendingSurat()
	}

	request.session['status']=""
	return render(request,'master/display_jabatan.html',context)

def delGolongan(request,id):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		MasterGolongan.objects.filter(kode_golongan=id).delete()
		addLogging(request.user.username,"master_golongan",f"berhasil - delete kode: {id}" )
		request.session['status']="Master Golongan Berhasil dihapus!"
	except:
		addLogging(request.user.username,"master_golongan",f"Gagal[data sudah terpakai] - delete kode: {id}" )
		request.session['status']="Master Golongan Gagal dihapus!"

	return HttpResponseRedirect('/master/gol/dis/')

def delJabatan(request,id):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		MasterJabatan.objects.filter(kode_jabatan=id).delete()
		addLogging(request.user.username,"master_jabatan",f"berhasil-delete kode: {id}" )
		request.session['status']="Master Golongan Berhasil dihapus!"
	except:
		addLogging(request.user.username,"master_jabatan",f"gagal[data sudah terpakai] - delete kode: {id}" )
		request.session['status']="Master Golongan Gagal dihapus!"
	return HttpResponseRedirect('/master/jab/dis/')

def addPegawai(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	if request.method=="POST":
		nik = request.POST['nik']
		nama = request.POST['nama']
		kode_jabatan = request.POST['kode_jabatan']
		kode_golongan = request.POST['kode_golongan']
		try:
			is_kepala = request.POST['is_kepala']
			is_kepala=True
		except:
			is_kepala=False

		try:
			pegawai = MasterPegawai.objects.create(
				nik = nik,
				nama=nama,
				kode_jabatan=MasterJabatan.objects.get(kode_jabatan=kode_jabatan),
				kode_golongan=MasterGolongan.objects.get(kode_golongan=kode_golongan),
				is_kepala=is_kepala,
				updatedBy=request.user.username
			)
			pegawai.save()
			print(pegawai)
			
			if(pegawai!=None):
				addLogging(request.user.username,"master_pegawai",f"berhasil-create NIP: {nik} - {nama}" )
				request.session['status']="Master Pegawai Berhasil dibuat!"
			else:
				addLogging(request.user.username,"master_pegawai",f"gagal[kepala pegawai sudah dibuat]-create NIP: {nik} - {nama}" )
				request.session['status']="Master Pegawai Gagal dibuat! Kepala Pegawai sudah pernah ada!"
		except:
			addLogging(request.user.username,"master_pegawai",f"gagal[master pegawai sudah ada]-create NIP: {nik} - {nama}" )
			request.session['status']="Master Pegawai Gagal dibuat! Master Pegawai sudah pernah ada atau Kepala Pegawai sudah ada!"
		return HttpResponseRedirect('/master/peg/dis/')
	
	context={
		'forms':inputPegawai,
		'menuname':'Menambah Pegawai',
		'pathway':'Master Pegawai - Menambah Pegawai',
		'username':request.user.username,
		'pending_surat':getPendingSurat()
	}
	return render(request,'master/create_pegawai.html',context)

def displayPegawai(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		page=request.GET['page']
	except:
		page=1
	if page==None:
		page=1

	listPegawai = MasterPegawai.objects.all().order_by('-updatedAt')
	
	p=Paginator(listPegawai,5)
	
	
	try:
		halaman = p.page(page)
	except:
		halaman = p.page(1)
	
	context={
		'lists':halaman,
		'paginator':p,
		'menuname':'Daftar Pegawai',
		'pathway':'Master Pegawai - Daftar Pegawai',
		'username':request.user.username,
		'mywebsite': '/master/peg/dis/',
		'status':request.session['status'],
		'pending_surat':getPendingSurat()
	}
	request.session['status']=""
	return render(request,'master/display_pegawai.html',context)

def delPegawai(request,id):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		MasterPegawai.objects.filter(nik=id).delete()
		addLogging(request.user.username,"master_pegawai",f"berhasil-delete kode: {id}" )
		request.session['status']="Master Pegawai Berhasil dihapus"
	except:
		addLogging(request.user.username,"master_pegawai",f"gagal[data sudah terpakai]-delete kode: {id}" )
		request.session['status']="Master Pegawai Gagal dihapus"
	return HttpResponseRedirect('/master/peg/dis/')

def addPengguna(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	
	if(request.user.is_superuser!=True):
		return HttpResponseRedirect('/')
	
	if request.method=="POST":
		username = request.POST['username']
		password = request.POST['password']
		email = request.POST['username']
		
		try:
			is_create = request.POST['is_create']
			is_create = True
		except:
			is_create=False

		try:
			is_edit = request.POST['is_edit']
			is_edit = True
		except:
			is_edit=False

		try:
			is_delete = request.POST['is_delete']
			is_delete = True
		except:
			is_delete=False

		try:
			pengguna = Pengguna.objects.create(
				username=username,
				password=password,
				email=email,
				is_create=is_create,
				is_edit=is_edit,
				is_delete=is_delete,
				updatedBy=request.user.username
			)
			pengguna.save()
			pengguna=Pengguna.objects.get(username=username)
			pengguna.set_password(password)
			pengguna.save()
			addLogging(request.user.username,"master_pengguna",f"berhasil-create username: {username} - {email}" )
			request.session['status']="Master Pengguna Berhasil ditambahkan!"
		except:
			addLogging(request.user.username,"master_pengguna",f"gagal[data sudah ada]-create username: {username} - {email}" )
			request.session['status']="Master Pengguna Gagal ditambahkan!"
		return HttpResponseRedirect('/master/user/dis/')
	
	context={
		'forms':inputPengguna,
		'menuname':'Menambah Pengguna Aplikasi',
		'pathway':'Master Pengguna - Menambah Pengguna',
		'username':request.user.username,
		'pending_surat':getPendingSurat()
	}
	return render(request,'master/create_pengguna.html',context)

def displayPengguna(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	
	if(request.user.is_superuser!=True):
		return HttpResponseRedirect('/')
	
	try:
		page=request.GET['page']
	except:
		page=1
	if page==None:
		page=1

	listPengguna = Pengguna.objects.all().order_by('-updatedAt')
	
	p=Paginator(listPengguna,5)
	
	
	try:
		halaman = p.page(page)
	except:
		halaman = p.page(1)
	
	context={
		'lists':halaman,
		'paginator':p,
		'menuname':'Daftar Pengguna Aplikasi',
		'pathway':'Master Pengguna - Daftar Pengguna',
		'username':request.user.username,
		'mywebsite': '/master/user/dis/',
		'status':request.session['status'],
		'pending_surat':getPendingSurat()
	}
	request.session['status']=""
	return render(request,'master/display_pengguna.html',context)

def delPengguna(request,id):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	
	if(request.user.is_superuser!=True):
		return HttpResponseRedirect('/')
	
	try:
		Pengguna.objects.filter(username=id).delete()
		addLogging(request.user.username,"master_pengguna",f"berhasil-delete kode: {id}" )
		request.session['status']="Master Pengguna Berhasil dihapus!"
	except:
		addLogging(request.user.username,"master_pengguna",f"gagal[data sudah dipakai]-delete kode: {id}" )
		request.session['status']="Master Pengguna Gagal dihapus!"
	return HttpResponseRedirect('/master/user/dis/')

def loginuser(request):
	if request.method=="POST":
		request.session['status']=''
		username=request.POST['username']
		password=request.POST['password']
		user = authenticate(username=username,password=password)
		if(user != None):
			login(request,user)
			if(request.user.is_authenticated):
				addLogging(request.user.username,"login",f"berhasil" )
			return HttpResponseRedirect('/')
		else:
			return HttpResponseRedirect('/auth/')
		
	if(request.user.is_authenticated):
		return HttpResponseRedirect('/')
	else:
		return render(request,'login.html')


def logoutuser(request):
	addLogging(request.user.username,"logout",f"berhasil" )
	logout(request)
	return HttpResponseRedirect('/auth/')

def displayLog(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	if(request.user.is_superuser !=True):
		return HttpResponseRedirect('/')
	
	#get the parameters query 
	print(request.GET.dict())

	try:
		page=request.GET['page']
	except:
		page=1
	if page==None:
		page=1

	listLogging = Logging.objects.all().order_by('-updatedAt')
	
	p=Paginator(listLogging,10)
	
	
	try:
		halaman = p.page(page)
	except:
		halaman = p.page(1)
	
	context={
		'lists':halaman,
		'paginator':p,
		'menuname':'Daftar Log User',
		'pathway':'log transaksi',
		'username':request.user.username,
		'mywebsite': '/logs/',
		'pending_surat':getPendingSurat()
	}
	return render(request,'master/display_logging.html',context)

def displayGolonganID(request,id):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		page=request.GET['page']
	except:
		page=1
	if page==None:
		page=1
	listGolongan = MasterGolongan.objects.all().filter(Q(kode_golongan__icontains=id) | Q(keterangan__icontains=id)).order_by('-updatedAt')
	p=Paginator(listGolongan,5)
	
	try:
		halaman = p.page(page)
	except:
		halaman = p.page(1)
	context={
		'lists':halaman,
		'paginator':p,
		'menuname':'Daftar Golongan  [' + str(id).upper() + "]",
		'pathway':'Master Golongan - Daftar Golongan',
		'username':request.user.username,
		'mywebsite': '/master/gol/dis/',
		'pending_surat':getPendingSurat()
	}
	return render(request,'master/display_golongan.html',context)

def displayJabatanID(request,id):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		page=request.GET['page']
	except:
		page=1
	if page==None:
		page=1
	listJabatan = MasterJabatan.objects.all().filter(Q(kode_jabatan__icontains=id) | Q(keterangan__icontains=id)).order_by('-updatedAt')
	p=Paginator(listJabatan,5)
	
	try:
		halaman = p.page(page)
	except:
		halaman = p.page(1)
	context={
		'lists':halaman,
		'paginator':p,
		'menuname':'Daftar Jabatan  [' + str(id).upper() + "]",
		'pathway':'Master Jabatan - Daftar Jabatan',
		'username':request.user.username,
		'mywebsite': '/master/jab/dis/',
		'pending_surat':getPendingSurat()
	}
	return render(request,'master/display_jabatan.html',context)

def displayPegawaiID(request,id):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		page=request.GET['page']
	except:
		page=1
	if page==None:
		page=1
	listPegawai = MasterPegawai.objects.all().filter(Q(nik__icontains=id) | Q(nama__icontains=id)).order_by('-updatedAt')
	p=Paginator(listPegawai,5)
	
	try:
		halaman = p.page(page)
	except:
		halaman = p.page(1)
	context={
		'lists':halaman,
		'paginator':p,
		'menuname':'Daftar Pegawai  [' + str(id).upper() + "]",
		'pathway':'Master Pegawai - Daftar Pegawai',
		'username':request.user.username,
		'mywebsite': '/master/peg/dis/',
		'pending_surat':getPendingSurat()
	}
	return render(request,'master/display_pegawai.html',context)

def displayPenggunaID(request,id):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	
	if(request.user.is_superuser!=True):
		return HttpResponseRedirect('/')

	try:
		page=request.GET['page']
	except:
		page=1
	if page==None:
		page=1
	listPengguna = Pengguna.objects.all().filter(Q(username__icontains=id) | Q(email__icontains=id)).order_by('-updatedAt')
	p=Paginator(listPengguna,5)
	
	try:
		halaman = p.page(page)
	except:
		halaman = p.page(1)
	context={
		'lists':halaman,
		'paginator':p,
		'menuname':'Daftar Pengguna Aplikasi  [' + str(id).upper() + "]",
		'pathway':'Master Pengguna - Daftar Pengguna',
		'username':request.user.username,
		'mywebsite': '/master/user/dis/',
		'pending_surat':getPendingSurat()
	}
	return render(request,'master/display_pengguna.html',context)

def displayLogID(request,id):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	if(request.user.is_superuser !=True):
		return HttpResponseRedirect('/')
	
	#get the parameters query 
	print(request.GET.dict())

	try:
		page=request.GET['page']
	except:
		page=1
	if page==None:
		page=1

	listLogging = Logging.objects.all().filter(Q(grouping__icontains=id) | Q(transaction__icontains=id)).order_by('-updatedAt')
	
	p=Paginator(listLogging,10)
	
	
	try:
		halaman = p.page(page)
	except:
		halaman = p.page(1)
	
	context={
		'lists':halaman,
		'paginator':p,
		'menuname':'Daftar Log User ['+str(id).upper()+']',
		'pathway':'log transaksi',
		'username':request.user.username,
		'mywebsite': '/logs/',
		'pending_surat':getPendingSurat()

	}
	return render(request,'master/display_logging.html',context)

def addNomorSurat(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	if request.method=="POST":
		nomor_surat = request.POST['nomor_surat']
		tgl_surat =  request.POST['tgl_surat']
		lokasi_surat = request.POST['lokasi_surat']
		tujuan = request.POST['tujuan']
		lokasi = request.POST['lokasi']
		tgl_awal_tugas = request.POST['tgl_awal_tugas']
		tgl_akhir_tugas = request.POST['tgl_akhir_tugas']

		confignya = ConfigDispenda.objects.all()

		if(confignya.count()==0):
			request.session['status']="Penambahan Surat Tugas! Kepala Kantor Belum Disetting!"
			addLogging(request.user.username,'trx_surat_tugas_header',"gagal[kepala kantor belum disetting]")
			return HttpResponseRedirect('/surat/dis/')
		else:
			try:
				id_surat=str(uuid.uuid4())
				trxsurattugas = TrxSuratTugas.objects.create(
					id_surat=id_surat,
					nomor_surat=nomor_surat,
					tgl_surat=tgl_surat,
					tgl_akhir_tugas = tgl_akhir_tugas,
					lokasi_surat = lokasi_surat,
					tujuan=tujuan,
					tgl_awal_tugas=tgl_awal_tugas,
					updatedBy=request.user.username,
					lokasi=lokasi,
					submit=False
				)
				
				request.session['status']="Penambahan Surat Tugas Berhasil!"
				addLogging(request.user.username,f'trx_surat_tugas_header',f"berhasil-Nomor Surat: {nomor_surat}")			
				return HttpResponseRedirect('/surat/add/'+id_surat)
			except:
				request.session['status']="Penambahan Surat Tugas Gagal! Nomor Surat Pernah Dibuat!"
				addLogging(request.user.username,f'trx_surat_tugas_header',"gagal[Nomor Surat Pernah Dibuat]")

	context={
		'forms':inputSuratTugas,
		'menuname':'Menambah Surat Tugas',
		'pathway':'Transaksi Surat Tugas - Menambah Surat Tugas',
		'username':request.user.username,
		'status':request.session['status']
	}
	request.session['status']=""
	return render(request,'master/create_surat_tugas.html',context)


def delNomorSurat(request,id):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		id_surat=id
		data=TrxSuratTugas.objects.all().filter(Q(id_surat=id_surat)&Q(submit=False))
		if data.count()>0:
			nomorSurat=data[0].nomor_surat
			data.delete()
			ST_DasarTugas.objects.all().filter(id_surat=nomorSurat).delete()
			ST_Peserta.objects.all().filter(id_surat=nomorSurat).delete()
			addLogging(request.user.username,"header surat tugas",f"berhasil - delete surat tugas nomor: {nomorSurat}" )
		else:
			addLogging(request.user.username,"header surat tugas",f"gagal - delete surat tugas nomor: {nomorSurat}" )	
	except:
		addLogging(request.user.username,"header surat tugas",f"gagal - delete surat tugas karena sudah disubmit" )
	return HttpResponseRedirect('/surat/dis/')

def displaySuratTugas(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')

	try:
		page=request.GET['page']
	except:
		page=1
	if page==None:
		page=1

	listTrxSuratTugas = TrxSuratTugas.objects.all().order_by('-updatedAt')
	
	p=Paginator(listTrxSuratTugas,10)
	
	
	try:
		halaman = p.page(page)
	except:
		halaman = p.page(1)
	
	context={
		'lists':halaman,
		'paginator':p,
		'menuname':'Daftar Surat Tugas',
		'pathway':'Transaksi Surat Tugas - Daftar Surat Tugas',
		'username':request.user.username,
		'mywebsite': '/logs/',
		'status':request.session['status'],
		'pending_surat':getPendingSurat()
	}
	request.session['status']=''
	return render(request,'master/display_surat_tugas.html',context)

def addDasarSurat(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	if request.method=="POST":
		kode_surat = request.POST['kode_dasar']
		keterangan = request.POST['keterangan']

		try:
			dasarsurat = MasterDasarST.objects.create(
				kode_dasar=kode_surat,
				keterangan=keterangan,
				updatedBy=request.user.username
			)
			
			addLogging(request.user.username,"master_dasarsurat",f"berhasil - create kode: {kode_surat} - {keterangan}" )
			request.session['status']="Penambahan Master Dasar Surat Tugas Berhasil!"
		except Exception as ex:
			# print(ex)
			request.session['status']="Penambahan Master Dasar Surat Tugas Gagal! Data pernah dibuat!"
			addLogging(request.user.username,"master_dasarsurat",f"gagal[data sudah ada] - create kode: {kode_surat} - {keterangan}" )
		return HttpResponseRedirect('/master/sur/dis/')

	context={
		'forms':inputMasterDasarST,
		'menuname':'Transaksi Surat Tugas',
		'pathway':'Master Surat Tugas - Menambah Surat Tugas',
		'username':request.user.username
	}
	request.session['status']=""
	return render(request,'master/create_dasarst.html',context)

def displayMasterDasarSurat(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		page=request.GET['page']
	except:
		page=1
	if page==None:
		page=1

	listSurat = MasterDasarST.objects.all().order_by('-updatedAt')
	
	p=Paginator(listSurat,5)
	
	
	try:
		halaman = p.page(page)
	except:
		halaman = p.page(1)
	
	context={
		'lists':halaman,
		'paginator':p,
		'menuname':'Daftar Dasar Surat Tugas',
		'pathway':'Master Dasar Surat Tugas - Daftar Dasar Surat Tugas',
		'username':request.user.username,
		'mywebsite': '/master/sur/dis/',
		'status':request.session['status'],
		'pending_surat':getPendingSurat()
	}
	request.session['status']=""
	return render(request,'master/display_dasarst.html',context)

def displayMasterDasarSuratID(request,id):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	if(request.user.is_superuser !=True):
		return HttpResponseRedirect('/')
	
	#get the parameters query 
	print(request.GET.dict())

	try:
		page=request.GET['page']
	except:
		page=1
	if page==None:
		page=1

	listDasar = MasterDasarST.objects.all().filter(Q(kode_dasar__icontains=id) | Q(keterangan__icontains=id)).order_by('-updatedAt')
	
	p=Paginator(listDasar,10)
	
	
	try:
		halaman = p.page(page)
	except:
		halaman = p.page(1)
	
	context={
		'lists':halaman,
		'paginator':p,
		'menuname':'Daftar Dasar Surat Tugas ['+str(id).upper()+']',
		'pathway':'Master Dasar Surat Tugas - Display Dasar Surat Tugas',
		'username':request.user.username,
		'mywebsite': '/master/sur/dis/',
		'pending_surat':getPendingSurat()

	}
	request.session['status']=""
	return render(request,'master/display_dasarst.html',context)

def delMasterDasarSuratTugas(request,id):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		MasterDasarST.objects.filter(kode_dasar=id).delete()
		addLogging(request.user.username,"master_dasarsurat",f"berhasil-delete kode: {id}" )
		request.session['status']="Master Dasar Surat Tugas berhasil dihapus!"
	except:
		addLogging(request.user.username,"master_dasarsurat",f"gagal[data sudah dipakai]-delete kode: {id}" )
		request.session['status']="Master Dasar Surat Tugas Gagal dihapus!"
	return HttpResponseRedirect('/master/sur/dis/')

def updateConfig(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	if(request.user.is_superuser !=True):
		return HttpResponseRedirect('/')
	
	if request.method=="POST":
		kepala = request.POST['kepala']
		try:
			is_PLT = request.POST['is_PLT']
			is_PLT=True
		except:
			is_PLT=False

		try:
			configured = ConfigDispenda.objects.create(
				kepala=MasterPegawai.objects.get(nik=kepala),
				is_PLT=is_PLT
			)
			configured.save()
			addLogging(request.user.username,"config",f"berhasil - kepala: {kepala} - {is_PLT}" )
			request.session['status']="Konfigurasi Sistem Berhasil!"
		except Exception as ex:
			print(ex)
			# print(ex)
			request.session['status']="Konfigurasi Sistem Gagal!"
			addLogging(request.user.username,"config",f"gagal[{ex}]" )
		return HttpResponseRedirect('/')

	kepala = ConfigDispenda.objects.all()
	if(kepala.count()==0):
		kepala=None
	else:
		kepala=kepala[0]

	context={
		'forms':inputConfig,
		'menuname':'Konfigurasi Aplikasi',
		'pathway':'Profile - Setting Sistem',
		'username':request.user.username,
		'kepala':kepala,
		'pending_surat':getPendingSurat()
	}
	request.session['status']=""
	return render(request,'master/display_config.html',context)


def detailSuratTugas(request,id):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		data = TrxSuratTugas.objects.get(id_surat=id)

		list_surat_tugas=ST_DasarTugas.objects.all().filter(id_surat=data)
		list_peserta_tugas=ST_Peserta.objects.all().filter(id_surat=data)

		form_dasar = inputDasarSuratTugas()
		form_peserta = inputPesertaTugas()

		context={
			'form_dasar':form_dasar,
			'form_peserta':form_peserta,
			'list_surat':list_surat_tugas,
			'list_peserta':list_peserta_tugas,
			'menuname':'Transaksi Surat Tugas',
			'pathway':'Master Surat Tugas - Menambah Surat Tugas - Detail',
			'username':request.user.username,
			'id_surat':id,
			'data':data,
			'pending_surat':getPendingSurat()
		}
		request.session['status']=""
		return render(request,'master/create_surat_tugas_detail.html',context)
	except Exception as ex:
		print(ex)
		return HttpResponseRedirect('/surat/dis/')
	
def detailSuratTugas_add_surat(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		if request.method=="POST":
			id_surat = request.GET['id']
			dasar_tugas = request.POST['dasar_tugas']
			
			surat = ST_DasarTugas.objects.create(
				id_surat = TrxSuratTugas.objects.get(id_surat=id_surat).nomor_surat,
				dasar_tugas = MasterDasarST.objects.get(kode_dasar=dasar_tugas),
				updatedBy=request.user.username
			)
			# print(surat.pk)
			if surat.pk is not None:
				request.session['status']="Detail Dasar Surat Tugas Berhasil Ditambahkan!"
				addLogging(request.user.username,"detail surat tugas",f"berhasil - tambah id_surat: {id_surat} - dasar surat: {dasar_tugas}" )
			else:
				request.session['status']="Detail Dasar Surat Tugas Gagal Ditambahkan!"				
				addLogging(request.user.username,"detail surat tugas",f"gagal - tambah id_surat: {id_surat} - dasar surat: {dasar_tugas}" )
		return HttpResponseRedirect('/surat/add/' + id_surat)
	except:
		return HttpResponseRedirect('/surat/dis/')
	
def detailSuratTugas_add_pegawai(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		if request.method=="POST":
			id_surat = request.GET['id']
			peserta = request.POST['peserta']
			print(MasterPegawai.objects.get(nik=peserta))
			pesertanya = ST_Peserta.objects.create(
				id_surat = TrxSuratTugas.objects.get(id_surat=id_surat).nomor_surat,
				peserta = MasterPegawai.objects.get(nik=peserta),
				updatedBy=request.user.username
			)

			if pesertanya.pk is not None:
				request.session['status']="Detail Pegawai Berhasil Ditambahkan!"
				addLogging(request.user.username,"detail surat tugas",f"berhasil - tambah id_surat: {id_surat} - peserta: {peserta}" )
			else:
				request.session['status']="Detail Pegawai Gagal Ditambahkan!"				
				addLogging(request.user.username,"detail surat tugas",f"gagal - tambah id_surat: {id_surat} - peserta: {peserta}" )
			
		return HttpResponseRedirect('/surat/add/' + id_surat)
	except:
		return HttpResponseRedirect('/surat/dis/')

def detailSuratTugas_del_surat(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		id_surat = request.GET['id']
		dasar_tugas = request.GET['dasar_tugas']
		cek_submit = TrxSuratTugas.objects.get(id_surat=id_surat)
		if(cek_submit.submit!=True):
			ST_DasarTugas.objects.all().filter(Q(id_surat=TrxSuratTugas.objects.get(id_surat=id_surat).nomor_surat) & Q(dasar_tugas = MasterDasarST.objects.get(kode_dasar=dasar_tugas))).delete()
			addLogging(request.user.username,"detail surat tugas",f"berhasil - hapus id_surat: {id_surat} - dasar surat: {dasar_tugas}" )
			request.session['status']="Detail Dasar Surat Tugas Berhasil Dihapus!"
		else:
			request.session['status']="Detail Dasar Surat Tugas Gagal Dihapus!"
			addLogging(request.user.username,"detail surat tugas",f"gagal - hapus id_surat: {id_surat} - dasar surat: {dasar_tugas}" )
		return HttpResponseRedirect('/surat/add/' + id_surat)
	except:
		return HttpResponseRedirect('/surat/dis/')

def detailSuratTugas_del_pegawai(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		id_surat = request.GET['id']
		peserta = request.GET['peserta']
		cek_submit = TrxSuratTugas.objects.get(id_surat=id_surat)
		if(cek_submit.submit!=True):
			ST_Peserta.objects.filter(Q(id_surat = TrxSuratTugas.objects.get(id_surat=id_surat).nomor_surat) & Q(peserta = MasterPegawai.objects.get(nik=peserta))).delete()
			addLogging(request.user.username,"detail surat tugas",f"berhasil - hapus id_surat: {id_surat} - dasar surat: {peserta}" )
			request.session['status']="Detail Peserta Tugas Berhasil Dihapus!"
		else:
			addLogging(request.user.username,"detail surat tugas",f"gagal - hapus id_surat: {id_surat} - dasar surat: {peserta}" )
			request.session['status']="Detail Peserta Tugas Gagal Dihapus!"
		return HttpResponseRedirect('/surat/add/' + id_surat)
	except:
		return HttpResponseRedirect('/surat/dis/')

def detailSuratTugas_submit(request):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		id_surat = request.GET['id']
		TrxSuratTugas.objects.filter(id_surat=id_surat).update(submit=True)		
		addLogging(request.user.username,"detail surat tugas",f"sukses submit id_surat: {id_surat}" )
		request.session['status']="Surat Tugas Berhasil Dikirim!"
	except:
		request.session['status']="Detail Surat Tugas Gagal dikirim!"	
	return HttpResponseRedirect('/surat/dis/')

def Exportkan(request,id):
	if(request.user.is_authenticated != True):
		return HttpResponseRedirect('/auth/')
	try:
		surattugas = TrxSuratTugas.objects.get(Q(id_surat=id) & Q(submit=True))
		nomorsurat=surattugas.nomor_surat
	except Exception as ex:
		print(ex)
		nomorsurat = None
	try:
		if nomorsurat is not None:
			document = docx.Document()

			font = document.styles['Normal'].font
			font.name = "Time New Roman"
			font.size = Pt(12)

			sections = document.sections
			for section in sections:
				section.top_margin = Cm(2)
				section.left_margin = Cm(2)
				section.right_margin = Cm(2)
				section.bottom_margin = Cm(2)

			document.add_picture(os.path.join(settings.BASE_DIR,'static/img/kop_surat.png'),width=Cm(17), height=Cm(4))
			paragraph = document.add_paragraph()
			runner = paragraph.add_run("SURAT PERINTAH TUGAS")
			runner.bold=True
			runner.underline=True
			runner = paragraph.add_run("\nNomor: %s"%surattugas.nomor_surat)
			paragraph.alignment = 1

			table = document.add_table(rows=1,cols=2)
			row = table.rows[0].cells
			paragraph=row[0].add_paragraph()
			runner = paragraph.add_run("DASAR:")
			runner.bold=True
			row[1].add_paragraph("Keputusan Menteri Dalam Negeri Nomor 16 Tahun 2013 tanggal 23 Januari 2013 tentang Pelaksanaan Perjalanan Dinas;",style="List Number")
			row[1].add_paragraph("Peraturan Daerah Provinsi Jawa Tengah Nomor 12 Tahun 2021 tentang Anggaran Pendapatan dan Belanja Daerah Provinsi Jawa Tengah Tahun Anggaran 2022;",style="List Number")
			row[1].add_paragraph("Peraturan Gubernur Jawa Tengah Nomor 17 Tahun 2013 tentang Perjalanan Dinas Gubernur / Wakil Gubernur, Pimpinan dan Anggota Dewan Perwakilan Rakyat Daerah, Pegawai Negeri Sipil, Calon Pegawai Negeri Sipil Dan Pegawai Non Pegawai Negeri Sipil;",style="List Number")
			row[1].add_paragraph("Peraturan Gubernur Jawa Tengah Nomor 27 Tahun 2020  tentang Standar Harga Satuan Provinsi Jawa;",style="List Number")
			
			#lookup for dasar tambahan
			dasar = ST_DasarTugas.objects.all().filter(id_surat=surattugas.nomor_surat)
			for dasarnya in dasar:
				row[1].add_paragraph(dasarnya.dasar_tugas.keterangan + ";",style="List Number")

			for cell in table.columns[0].cells:
				cell.width = Cm(3)
			for cell in table.columns[1].cells:
				cell.width = Cm(14)

			paragraph = document.add_paragraph()
			runner = paragraph.add_run("\nMEMERINTAHKAN")
			runner.bold=True
			paragraph.alignment = 1

			table = document.add_table(rows=1,cols=2)
			row = table.rows[0].cells

			peserta_tugas = ST_Peserta.objects.all().filter(id_surat=surattugas.nomor_surat)
			
			paragraph = row[0].add_paragraph()
			runner = paragraph.add_run("KEPADA:")
			runner.bold=True
			parnya = row[1].add_paragraph("1.   Nama:\t" + peserta_tugas[0].peserta.nama)
			runner = parnya.add_run("\n      NIP:\t" + peserta_tugas[0].peserta.nik)
			runner = parnya.add_run("\n      Jabatan:\t" + peserta_tugas[0].peserta.kode_jabatan.keterangan)


			#lookup for peserta tambahan
			counter=2
			peserta_ke=0
			for pesertanya in peserta_tugas:
				if peserta_ke>0:
					parnya = row[1].add_paragraph(str(counter)+".   Nama:\t" + pesertanya.peserta.nama)
					runner = parnya.add_run("\n      NIP:\t" + pesertanya.peserta.nik)
					runner = parnya.add_run("\n      Jabatan:\t" + pesertanya.peserta.kode_jabatan.keterangan)
				peserta_ke+=1

			for cell in table.columns[0].cells:
				cell.width = Cm(3)
			for cell in table.columns[1].cells:
				cell.width = Cm(14)

			table = document.add_table(rows=1,cols=2)
			row = table.rows[0].cells
			paragraph=row[0].add_paragraph()
			runner = paragraph.add_run("UNTUK:")
			runner.bold=True
			mypar=row[1].add_paragraph("")
			runner = mypar.add_run("1.   Melaksanakan tugas:\t%s\n"%surattugas.tujuan)
			runner = mypar.add_run("      di:\t\t\t\t%s\n"%surattugas.lokasi)
			runner = mypar.add_run("      tanggal:\t\t\t%s s/d %s\n"%(surattugas.tgl_awal_tugas,surattugas.tgl_akhir_tugas))
			runner = mypar.add_run("2.   Tidak menerima gratifikasi dalam bentuk apapun sesuai ketentuan;\n")
			runner = mypar.add_run("3.   Melapor kepada Pejabat setempat guna pelaksanaan tugas tersebut;\n")
			runner = mypar.add_run("4.   Melaporkan Hasil Pelaksanaan Tugas kepada Pejabat pemberi tugas.")
			print(surattugas)
			
			for cell in table.columns[0].cells:
				cell.width = Cm(3)
			for cell in table.columns[1].cells:
				cell.width = Cm(14)

			paragraph = document.add_paragraph()
			runner = paragraph.add_run("\n\t\t\t\t\t\tDitetapkan di:\tSemarang\n")
			runner = paragraph.add_run("\t\t\t\t\t\tpada tanggal:\t\t%s\n" %surattugas.tgl_surat)
			runner = paragraph.add_run("\t\t\t\t\t\t___________________________________________________\n")
			
			Config = ConfigDispenda.objects.all()[0]
			if(Config.is_PLT):
				PLT="Plt. "
			else:
				PLT=""
			runner = paragraph.add_run(f"\t\t\t\t\t\t{PLT}KEPADA BADAN PENGELOLA PENDAPATAN\n")
			runner = paragraph.add_run("\t\t\t\t\t\tDAERAH PROVINSI JAWA TENGAH\n")
			runner = paragraph.add_run("\t\t\t\t\t\t"+Config.kepala.kode_jabatan.keterangan)
			runner = paragraph.add_run("\n\n\n\n\n\n")
			runner = paragraph.add_run("\t\t\t\t\t\t")
			runner = paragraph.add_run(Config.kepala.nama)
			runner.underline=True
			runner = paragraph.add_run("\n\t\t\t\t\t\t"+Config.kepala.kode_golongan.keterangan)

			document.save(os.path.join(settings.BASE_DIR,r"export\\" + str(id) + '.docx'))

			document = open(os.path.join(settings.BASE_DIR,r"export\\" + str(id) + '.docx'),'rb')
	
			resp = HttpResponse(document,content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
			filename = str(id) + '.docx'
			resp['Content-Disposition']='attachment;filename=%s'%format(filename)
			# request.session['status']="Surat Tugas Berhasil Di Download (Unduh)!"
			return resp
			# return HttpResponseRedirect('/surat/dis/')	
	except Exception as ex:
		print(ex)
		# request.session['status']="Surat Tugas Gagal Di Download (Unduh)!"
		return HttpResponseRedirect('/surat/dis/')		
	else:
		return HttpResponseRedirect('/surat/dis/')